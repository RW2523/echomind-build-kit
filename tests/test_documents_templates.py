"""Document templates — the `build_*` composers in server/mcp/documents.py.

These tests pass dicts straight in and read the Markdown straight out: the builders are
pure by design, so nothing here needs the seeded database, and a template that started
querying would fail this file before it failed a reviewer.

The load-bearing test is the last one. Every other test says "the value I passed came
out"; `test_no_builder_prints_a_number_absent_from_its_input` says the reverse — that no
number came out which did not go in, apart from figures the document itself names as
derived. That is golden rule 1 written as an assertion rather than as a hope.
"""

from __future__ import annotations

import io
import json
import re
from decimal import Decimal

import pytest

from server.mcp import documents as D

# --- fixtures as plain data --------------------------------------------------------

INVOICE_LINES = [
    {"instrument": "Confocal C2", "description": "Instrument time, peak",
     "qty": Decimal("6.00"), "amount": Decimal("289.00")},
    {"instrument": "Orbitrap Exploris", "description": "LC-MS/MS run",
     "hours": Decimal("4.50"), "amount": Decimal("612.50")},
    {"instrument": None, "description": "Consumables", "amount": Decimal("88.25")},
]
INVOICE_TOTAL = Decimal("989.75")  # == sum of the three lines above

BOOKING = {
    "id": "bk-4417",
    "instrument": "Titan Krios",
    "facility": "Advanced Imaging Core",
    "campus": "North Campus",
    "building": "Wellcome Building",
    "room": "B-014",
    "starts_at": "2026-09-03T09:00:00+00:00",
    "ends_at": "2026-09-03T13:30:00+00:00",
    "status": "requested",
    "account_code": "ACC-A1",
}

FACILITIES = [
    {
        "id": "fac-imaging", "name": "Advanced Imaging Core", "code": "IMG",
        "campus": "North Campus", "building": "Wellcome Building", "room": "B-014",
        "address": "Gower Place, London", "contact_email": "imaging@example.ac.uk",
        "opening_hours": "Mon-Fri 08:00-20:00",
        "instruments": [
            {"id": "ins-em-titan", "name": "Titan Krios",
             "techniques": ["cryo-EM", "single particle analysis"],
             "hourly_rate": Decimal("310.00"), "status": "available"},
            {"id": "ins-lightsheet", "name": "Light Sheet Z1",
             "techniques": ["light sheet microscopy"],
             "hourly_rate": Decimal("95.00"), "status": "maintenance"},
        ],
    },
    {
        "id": "fac-massspec", "name": "Mass Spectrometry Core", "code": "MSC",
        "campus": "Riverside Campus", "building": "Perutz Labs", "room": "L-002",
        "address": "Riverside Way, London", "contact_email": "massspec@example.ac.uk",
        "opening_hours": "Mon-Fri 09:00-17:00",
        "instruments": [],
    },
]

MATCHES = [
    {
        "id": "ins-em-titan", "name": "Titan Krios", "facility": "Advanced Imaging Core",
        "campus": "North Campus", "building": "Wellcome Building", "room": "B-014",
        "matched_on": ["cryo-EM"], "modality": "electron microscopy",
        "techniques": ["cryo-EM", "single particle analysis"],
        "sample_types": ["vitrified grids"], "specification": "300 kV, K3 detector",
        "hourly_rate": Decimal("310.00"), "status": "available",
    },
    {
        "id": "ins-nanopore", "name": "PromethION", "facility": "Genomics Core",
        "campus": "North Campus", "building": "Crick Wing", "room": "G-101",
        "match_reason": "long-read sequencing is the closest listed technique",
        "techniques": ["long-read sequencing"], "sample_types": ["gDNA"],
        "hourly_rate": Decimal("45.00"), "status": "available",
    },
]

USAGE_ROWS = [
    {"instrument": "Confocal C2", "month": "2026-03",
     "scheduled_hours": Decimal("12.00"), "tracked_hours": Decimal("10.50")},
    {"instrument": "MiSeq M3", "month": "2026-03",
     "scheduled_hours": Decimal("4.00"), "tracked_hours": Decimal("4.00")},
]
USAGE_TOTALS = {"scheduled_hours": Decimal("16.00"), "tracked_hours": Decimal("14.50"),
                "difference_hours": Decimal("-1.50")}


def _every_document() -> dict[str, str]:
    """One of everything, so the render and no-invention tests cover all five builders."""
    return {
        "invoice": D.build_invoice_statement(
            "ACC-A1", "2026-03", INVOICE_LINES, INVOICE_TOTAL,
            {"id": "lab-a", "name": "Lab A"},
        ),
        "booking": D.build_booking_confirmation(BOOKING),
        "directory": D.build_facility_directory(FACILITIES),
        "capability": D.build_capability_report("cryo-EM of vitrified grids", MATCHES),
        "usage": D.build_usage_summary(
            {"id": "u-alice", "name": "Alice Chen"}, "2026-03", USAGE_ROWS, USAGE_TOTALS,
        ),
    }


# --- invoice statement -------------------------------------------------------------


def test_an_invoice_statement_lists_every_line_with_its_hours_and_amount():
    md = D.build_invoice_statement("ACC-A1", "2026-03", INVOICE_LINES, INVOICE_TOTAL,
                                   {"id": "lab-a", "name": "Lab A"})
    assert md.startswith("# Invoice statement — ACC-A1")
    assert "| Instrument | Description | Hours | Amount |" in md
    assert "| Confocal C2 | Instrument time, peak | 6.00 | $289.00 |" in md
    assert "| Orbitrap Exploris | LC-MS/MS run | 4.50 | $612.50 |" in md
    assert "Lab A (`lab-a`)" in md
    assert "Charge lines: 3" in md


def test_an_invoice_total_equals_the_sum_of_its_lines():
    md = D.build_invoice_statement("ACC-A1", "2026-03", INVOICE_LINES, INVOICE_TOTAL, "Lab A")
    assert sum(line["amount"] for line in INVOICE_LINES) == INVOICE_TOTAL
    assert "**Total for 2026-03: $989.75**" in md
    assert "does not match the invoice total" not in md


def test_an_invoice_total_that_disagrees_with_its_lines_is_flagged():
    md = D.build_invoice_statement("ACC-A1", "2026-03", INVOICE_LINES, Decimal("1200.00"),
                                   "Lab A")
    assert "**Total for 2026-03: $1200.00**" in md
    assert "The lines above sum to $989.75" in md
    assert "does not match the invoice total" in md


def test_an_invoice_total_derived_from_the_lines_says_it_was_derived():
    md = D.build_invoice_statement("ACC-A1", "2026-03", INVOICE_LINES, None, "Lab A")
    assert "**Total for 2026-03: $989.75**" in md
    assert "derived by adding the lines above" in md


def test_an_invoice_statement_cites_the_fourteen_day_review_window():
    md = D.build_invoice_statement("ACC-A1", "2026-03", INVOICE_LINES, INVOICE_TOTAL, "Lab A")
    assert f"review period is {D.INVOICE_REVIEW_DAYS} days" in md
    assert D.INVOICE_REVIEW_DAYS == 14
    assert "silence is agreement" in md


def test_an_invoice_with_no_lines_refuses_to_show_a_total_it_cannot_derive():
    md = D.build_invoice_statement("ACC-A1", "2026-03", [], None, "Lab A")
    assert "_no charges in this period_" in md
    # Zero lines sums to zero, and "$0.00" would tell the reader they owe nothing —
    # a different claim from "no lines were supplied".
    assert "**Total for 2026-03: not recorded**" in md
    assert "$0.00" not in md


# --- booking confirmation ----------------------------------------------------------


def test_a_booking_confirmation_names_the_instrument_place_and_window():
    md = D.build_booking_confirmation(BOOKING)
    assert md.startswith("# Booking confirmation — Titan Krios")
    assert "Reference: `bk-4417`" in md
    assert "Facility: Advanced Imaging Core" in md
    assert "Where: North Campus · Wellcome Building, room B-014" in md
    assert "Starts: 2026-09-03T09:00:00+00:00" in md
    assert "Ends: 2026-09-03T13:30:00+00:00" in md
    assert "Account code: `ACC-A1`" in md
    assert "Status: **requested**" in md


def test_a_booking_duration_is_derived_from_the_window_and_labelled_as_such():
    md = D.build_booking_confirmation(BOOKING)
    assert "Duration: 4.50 h — derived from the window above" in md


def test_a_booking_confirmation_says_how_to_cancel_it():
    md = D.build_booking_confirmation(BOOKING)
    assert "## Changing or cancelling this booking" in md
    assert "`bk-4417`" in md.split("## Changing or cancelling this booking")[1]
    assert "The calendar is the record." in md
    assert "Booking and Cancellation Rules" in md


def test_a_booking_confirmation_omits_a_field_it_was_not_given():
    md = D.build_booking_confirmation(
        {"id": "bk-1", "instrument": "MiSeq M3", "facility": "Genomics Core",
         "starts_at": "2026-09-03T09:00:00+00:00", "ends_at": "2026-09-03T10:00:00+00:00",
         "status": "confirmed"}
    )
    assert "Account code: not recorded" in md
    assert "Where: not recorded" in md
    assert "ACC-" not in md


def test_a_booking_with_an_unusable_window_states_no_duration():
    md = D.build_booking_confirmation({**BOOKING, "ends_at": None})
    assert "Duration: not recorded" in md
    assert "derived from the window" not in md


def test_a_booking_prefers_the_instruments_own_room_over_the_facilitys():
    """Campus and building come from the core; the room is the door you knock on."""
    md = D.build_booking_confirmation({
        "id": "bk-9", "instrument": "Orbitrap Exploris",
        "facility": {"id": "fac-massspec", "name": "Mass Spectrometry Core",
                     "campus": "Riverside Campus", "building": "Perutz Labs",
                     "room": "L-002"},
        "room": "L-114",
        "starts_at": "2026-09-04T09:00:00+00:00", "ends_at": "2026-09-04T11:00:00+00:00",
        "status": "confirmed", "account_code": "ACC-B2",
    })
    assert "Facility: Mass Spectrometry Core (`fac-massspec`)" in md
    assert "Where: Riverside Campus · Perutz Labs, room L-114" in md
    assert "Duration: 2.00 h — derived from the window above" in md


# --- facility directory ------------------------------------------------------------


def test_a_facility_directory_gives_each_core_its_location_and_instruments():
    md = D.build_facility_directory(FACILITIES)
    assert "# Facility directory" in md
    assert "Cores listed: 2" in md
    assert "## Advanced Imaging Core" in md
    assert "Code: `IMG`" in md
    assert "Campus: North Campus" in md
    assert "Building: Wellcome Building" in md
    assert "Address: Gower Place, London" in md
    assert "Contact: imaging@example.ac.uk" in md
    assert "Opening hours: Mon-Fri 08:00-20:00" in md
    assert "| Titan Krios | cryo-EM, single particle analysis | $310.00 | available |" in md
    assert "| Light Sheet Z1 | light sheet microscopy | $95.00 | maintenance |" in md


def test_a_core_with_no_instruments_says_so_rather_than_borrowing_one():
    md = D.build_facility_directory(FACILITIES)
    massspec = md.split("## Mass Spectrometry Core")[1]
    assert "_no instruments listed for this core_" in massspec
    assert "Titan Krios" not in massspec


def test_an_empty_facility_directory_lists_nothing():
    md = D.build_facility_directory([])
    assert "Cores listed: 0" in md
    assert "No facilities were supplied" in md


# --- capability report -------------------------------------------------------------


def test_a_capability_report_states_the_goal_and_why_each_instrument_matched():
    md = D.build_capability_report("cryo-EM of vitrified grids", MATCHES)
    assert md.startswith("# Capability report — cryo-EM of vitrified grids")
    assert "Asked for: cryo-EM of vitrified grids" in md
    assert "Instruments matched: 2" in md
    assert "Matched on: cryo-EM" in md
    assert "Matched on: long-read sequencing is the closest listed technique" in md
    assert "### 1. Titan Krios" in md
    assert "### 2. PromethION" in md


def test_a_capability_report_gives_rate_status_and_where_the_instrument_is():
    md = D.build_capability_report("cryo-EM of vitrified grids", MATCHES)
    assert "Rate: $310.00 per hour" in md
    assert "Status: **available**" in md
    assert "Where: North Campus · Wellcome Building, room B-014" in md
    assert "Where: North Campus · Crick Wing, room G-101" in md
    assert "Specification: 300 kV, K3 detector" in md


def test_a_capability_report_with_no_matches_recommends_nothing():
    md = D.build_capability_report("neutron diffraction", [])
    assert "Instruments matched: 0" in md
    assert "## No match" in md
    assert "Nothing in the catalogue matched this request." in md
    for row in MATCHES:
        assert row["name"] not in md


# --- usage summary -----------------------------------------------------------------


def test_a_usage_summary_shows_scheduled_tracked_and_the_derived_difference():
    md = D.build_usage_summary({"id": "u-alice", "name": "Alice Chen"}, "2026-03",
                               USAGE_ROWS, USAGE_TOTALS)
    assert md.startswith("# Usage summary — Alice Chen (`u-alice`)")
    assert "| Instrument | Month | Scheduled h | Tracked h | Difference h (derived) |" in md
    assert "| Confocal C2 | 2026-03 | 12.00 | 10.50 | -1.50 |" in md
    assert "| MiSeq M3 | 2026-03 | 4.00 | 4.00 | 0.00 |" in md
    assert "| **Total** | | 16.00 | 14.50 | -1.50 |" in md
    assert "tracked hours minus scheduled hours, derived per row" in md


def test_a_usage_total_matches_the_sum_of_its_rows():
    md = D.build_usage_summary("lab-a", "2026-03", USAGE_ROWS, USAGE_TOTALS)
    assert sum(r["scheduled_hours"] for r in USAGE_ROWS) == USAGE_TOTALS["scheduled_hours"]
    assert sum(r["tracked_hours"] for r in USAGE_ROWS) == USAGE_TOTALS["tracked_hours"]
    assert "does not match the rows above it" not in md


def test_usage_totals_that_disagree_with_the_rows_are_flagged():
    md = D.build_usage_summary("lab-a", "2026-03", USAGE_ROWS,
                               {"scheduled_hours": Decimal("99.00"),
                                "tracked_hours": Decimal("14.50")})
    assert "does not match the rows above it" in md
    assert "scheduled was given as 99.00 against 16.00" in md
    # The tracked total agrees with its rows, so it is not dragged into the complaint.
    assert "tracked was given as" not in md


def test_a_supplied_difference_that_contradicts_its_own_totals_is_flagged():
    md = D.build_usage_summary("lab-a", "2026-03", USAGE_ROWS,
                               {**USAGE_TOTALS, "difference_hours": Decimal("0.00")})
    assert "difference was given as 0.00 against -1.50" in md


def test_a_usage_summary_without_totals_derives_them_from_the_rows():
    md = D.build_usage_summary("lab-a", None, USAGE_ROWS, None)
    assert "Period: all recorded months" in md
    assert "| **Total** | | 16.00 | 14.50 | -1.50 |" in md


def test_a_usage_summary_with_no_rows_says_so():
    md = D.build_usage_summary("lab-a", "2026-03", [], None)
    assert "_no usage recorded_" in md
    assert "Rows: 0" in md


# --- the guarantees ----------------------------------------------------------------


# A value from each document that a reader would notice the absence of — the point of
# rendering is that the numbers arrive, not that the call returned bytes.
LANDMARK = {
    "invoice": "989.75",
    "booking": "Titan Krios",
    "directory": "310.00",
    "capability": "Crick Wing",
    "usage": "14.50",
}


def _rendered_text(body: str, fmt: str) -> str:
    """What a reader of the rendered file would actually see."""
    blob = D.render("template test", body, fmt)
    assert isinstance(blob, bytes)
    if fmt == "md":
        return blob.decode("utf-8")
    if fmt == "docx":
        from docx import Document

        doc = Document(io.BytesIO(blob))
        cells = [c.text for table in doc.tables for row in table.rows for c in row.cells]
        return "\n".join([p.text for p in doc.paragraphs] + cells)

    from pypdf import PdfReader

    return "\n".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(blob)).pages)


@pytest.mark.parametrize("name", sorted(_every_document()))
@pytest.mark.parametrize("fmt", D.FORMATS)
def test_every_builder_survives_render_to_every_format(name, fmt):
    text = _rendered_text(_every_document()[name], fmt)
    assert LANDMARK[name] in text, f"{name} lost {LANDMARK[name]!r} on the way to {fmt}"


@pytest.mark.parametrize("fmt", D.FORMATS)
def test_a_rendered_table_keeps_its_rows(fmt):
    """The charge lines are the document; a renderer that drops one is a silent error."""
    text = _rendered_text(_every_document()["invoice"], fmt)
    for line in INVOICE_LINES:
        assert str(line["amount"]) in text


def test_a_pipe_inside_a_value_does_not_break_the_rendered_table():
    """A pipe in a description used to make a row wider than its header.

    Markdown tolerates that; reportlab refuses a ragged table outright, so the failure
    only ever appeared at PDF time, on a document the caller had already approved.
    """
    md = D.build_invoice_statement(
        "ACC-A1", "2026-03",
        [{"instrument": "Confocal C2", "description": "peak | out of hours",
          "hours": Decimal("2.00"), "amount": Decimal("100.00")}],
        Decimal("100.00"), "Lab A",
    )
    row = next(line for line in md.splitlines() if "Confocal C2" in line)
    assert row.count("|") == 5
    for fmt in D.FORMATS:
        assert D.render("pipes", md, fmt)


def test_a_newline_inside_a_value_does_not_end_the_table_early():
    md = D.build_facility_directory([
        {"name": "Genomics Core", "code": "GEN",
         "instruments": [{"name": "NovaSeq\nX Plus", "techniques": ["RNA-seq"],
                          "hourly_rate": Decimal("60.00"), "status": "available"}]},
    ])
    assert "| NovaSeq X Plus | RNA-seq | $60.00 | available |" in md
    for fmt in D.FORMATS:
        assert D.render("newlines", md, fmt)


def _numbers(text: str) -> set[str]:
    return set(re.findall(r"\d+(?:\.\d+)?", text))


# Figures the documents state they derived, plus the one policy constant they cite.
# Anything outside this set that is not in the input data is an invented number.
DERIVED = {
    "invoice": {"3"},              # the line count
    "booking": {"4.50"},           # the duration, derived from the window
    "directory": {"2"},            # the core count
    "capability": {"1", "2"},      # the match count and the numbered sections
    "usage": {"2", "0.00"},        # the row count and MiSeq's derived difference
}


@pytest.mark.parametrize("name", sorted(_every_document()))
def test_no_builder_prints_a_number_absent_from_its_input(name):
    """Golden rule 1, as arithmetic: no figure appears that was not passed or derived."""
    inputs = {
        "invoice": ["ACC-A1", "2026-03", INVOICE_LINES, INVOICE_TOTAL,
                    {"id": "lab-a", "name": "Lab A"}],
        "booking": [BOOKING],
        "directory": [FACILITIES],
        "capability": ["cryo-EM of vitrified grids", MATCHES],
        "usage": [{"id": "u-alice", "name": "Alice Chen"}, "2026-03", USAGE_ROWS,
                  USAGE_TOTALS],
    }[name]
    allowed = (
        _numbers(json.dumps(inputs, default=str))
        | DERIVED[name]
        | {str(D.INVOICE_REVIEW_DAYS)}
    )
    printed = _numbers(_every_document()[name])
    assert printed <= allowed, f"{name} invented {sorted(printed - allowed)}"


# --- the renderers, wired end to end (2026-08-12) ----------------------------------


@pytest.mark.tools
def test_every_registered_document_template_has_a_renderer():
    """A template the tool accepts but the executor cannot render is an approval that
    fails after the human has already said yes."""
    import inspect

    from server.mcp import actions
    from server.mcp import tools as T
    source = inspect.getsource(actions._exec_document)
    for template in T.DOCUMENT_TEMPLATES:
        assert f'"{template}"' in source, f"{template} has no renderer in _exec_document"


@pytest.mark.tools
def test_booking_confirmation_renders_from_a_real_booking():
    from sqlalchemy import text

    from server.db import session_scope
    from server.mcp import actions
    with session_scope() as s:
        bid = s.execute(
            text("SELECT id FROM infinity.bookings WHERE user_id='u-alice' LIMIT 1")
        ).scalar_one()
    title, body = actions._render_booking_confirmation("u-alice", {"booking_id": bid})
    assert bid in title and bid in body
    assert "Facility:" in body and "Where:" in body


@pytest.mark.tools
def test_a_booking_that_is_not_yours_is_not_found():
    """Scoped in SQL, so the answer is the same whether or not the booking exists."""
    from sqlalchemy import text

    from server.db import session_scope
    from server.mcp import actions
    from server.mcp.errors import ToolError
    with session_scope() as s:
        other = s.execute(
            text("SELECT id FROM infinity.bookings WHERE user_id='u-bob' LIMIT 1")
        ).scalar_one()
    for booking_id in (other, "bk-does-not-exist"):
        with pytest.raises(ToolError) as exc:
            actions._render_booking_confirmation("u-alice", {"booking_id": booking_id})
        assert exc.value.code == "not_found"


@pytest.mark.tools
def test_usage_summary_renders_through_the_scoped_tool():
    from server.mcp import actions
    title, body = actions._render_usage_summary("u-alice", {"month": "2026-03"})
    assert "u-alice" in title
    assert "Scheduled h" in body and "Tracked h" in body


@pytest.mark.tools
def test_the_facility_catalogue_exposes_location_and_capability():
    """get_facility_catalog predates migration 008 and returned id/name/code only, so the
    tool the planner reaches for most knew least about the facility."""
    from server.auth import Ctx
    from server.mcp import tools as T
    ctx = Ctx(user_id="u-alice", name="Alice", role="user", lab_ids=("lab-a",),
              facility_ids=(), raw={})
    out = T.get_facility_catalog(ctx)
    assert len(out["facilities"]) == 3 and len(out["instruments"]) == 12
    facility = next(f for f in out["facilities"] if f["id"] == "fac-imaging")
    for key in ("campus", "building", "room", "address", "opening_hours", "contact_email"):
        assert facility.get(key), f"{key} missing from the catalogue"
    titan = next(i for i in out["instruments"] if i["name"] == "Cryo-EM Titan")
    assert "cryo-EM" in titan["techniques"]
    assert titan["modality"] == "electron microscopy" and titan["room"]
