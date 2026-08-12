"""Render a generated document to DOCX or PDF.

The three templates in `actions.py` build Markdown, which was the whole output format:
a facility admin who asks for a monthly summary wants something they can attach to an
email, not a .md file. The renderers are unchanged — they still return Markdown, still
build every value from a query — and this turns that one representation into the two
formats people actually send.

Deliberately not a general Markdown engine. It handles exactly what the templates emit —
headings, pipe tables, bullets, `**bold**`, `inline code`, and the two-space line break —
and anything else passes through as plain text. A general parser would be a dependency
and a surface area for no gain, since the input is ours.

The second half of the file is the other end of the same pipe: the `build_*` templates
that compose the Markdown these renderers consume. They take rows a tool already fetched
and touch nothing else — no session, no clock, no config. A template that queries can
only be tested with a seeded database standing behind it, and the part worth testing is
the composition, not the query.
"""

from __future__ import annotations

import io
import re
from datetime import datetime
from decimal import Decimal, InvalidOperation

from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

FORMATS = ("md", "docx", "pdf")
EXTENSION = {"md": ".md", "docx": ".docx", "pdf": ".pdf"}

_HEADING_RE = re.compile(r"^(#{1,4})\s+(.*)$")
_BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
_TABLE_SEP_RE = re.compile(r"^\s*\|[\s:|-]+\|\s*$")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_CODE_RE = re.compile(r"`([^`]+)`")


def _cells(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _plain(text: str) -> str:
    """Strip the inline markers, for renderers that cannot show them."""
    return _CODE_RE.sub(r"\1", _BOLD_RE.sub(r"\1", text)).strip()


def _blocks(markdown: str) -> list[tuple[str, object]]:
    """Parse into ("heading"|"table"|"bullet"|"para", payload).

    One pass, because a pipe table has to be gathered across lines before it can be
    emitted, and everything else is line-at-a-time.
    """
    out: list[tuple[str, object]] = []
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        heading = _HEADING_RE.match(stripped)
        if heading:
            out.append(("heading", (len(heading.group(1)), _plain(heading.group(2)))))
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and _TABLE_SEP_RE.match(lines[i + 1]):
            header = _cells(stripped)
            rows: list[list[str]] = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([_plain(c) for c in _cells(lines[i])])
                i += 1
            out.append(("table", ([_plain(h) for h in header], rows)))
            continue

        bullet = _BULLET_RE.match(line)
        if bullet:
            out.append(("bullet", _plain(bullet.group(1))))
            i += 1
            continue

        out.append(("para", stripped))
        i += 1
    return out


def to_docx(title: str, markdown: str) -> bytes:
    doc = Document()
    doc.core_properties.title = title

    for kind, payload in _blocks(markdown):
        if kind == "heading":
            level, text = payload
            doc.add_heading(text, level=min(level, 4))
        elif kind == "bullet":
            doc.add_paragraph(str(payload), style="List Bullet")
        elif kind == "table":
            header, rows = payload
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for cell, text in zip(table.rows[0].cells, header, strict=True):
                cell.text = text
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for row in rows:
                cells = table.add_row().cells
                # Templates can emit a short row (the "no usage recorded" placeholder),
                # so pad rather than let zip(strict=True) raise on a document that is
                # otherwise perfectly valid.
                for cell, text in zip(cells, row + [""] * (len(header) - len(row)), strict=False):
                    cell.text = text
        else:
            para = doc.add_paragraph()
            _add_runs(para, str(payload))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_runs(paragraph, text: str) -> None:
    """Keep bold and code visibly distinct instead of flattening the markers away."""
    position = 0
    for match in re.finditer(r"\*\*(.+?)\*\*|`([^`]+)`", text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        if match.group(1) is not None:
            paragraph.add_run(match.group(1)).bold = True
        else:
            run = paragraph.add_run(match.group(2))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def to_pdf(title: str, markdown: str) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4, title=title,
        leftMargin=20 * mm, rightMargin=20 * mm, topMargin=18 * mm, bottomMargin=18 * mm,
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=9.5, leading=13)
    story: list[object] = []

    for kind, payload in _blocks(markdown):
        if kind == "heading":
            level, text = payload
            style = styles["Title"] if level == 1 else styles[f"Heading{min(level, 4)}"]
            story += [Paragraph(_escape(text), style), Spacer(1, 4)]
        elif kind == "bullet":
            story.append(Paragraph(f"• {_escape(str(payload))}", body))
        elif kind == "table":
            header, rows = payload
            data = [[Paragraph(f"<b>{_escape(h)}</b>", body) for h in header]]
            for row in rows:
                padded = row + [""] * (len(header) - len(row))
                data.append([Paragraph(_escape(c), body) for c in padded])
            table = Table(data, hAlign="LEFT", repeatRows=1)
            table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#c9d3dd")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef2f6")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story += [Spacer(1, 4), table, Spacer(1, 6)]
        else:
            story.append(Paragraph(_inline_to_html(str(payload)), body))

    doc.build(story)
    return buffer.getvalue()


def _escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _inline_to_html(text: str) -> str:
    """reportlab paragraphs take a small HTML dialect, so translate rather than strip."""
    escaped = _escape(text)
    escaped = _BOLD_RE.sub(r"<b>\1</b>", escaped)
    return _CODE_RE.sub(r'<font face="Courier">\1</font>', escaped)


def render(title: str, markdown: str, fmt: str) -> bytes:
    if fmt == "md":
        return markdown.encode("utf-8")
    if fmt == "docx":
        return to_docx(title, markdown)
    if fmt == "pdf":
        return to_pdf(title, markdown)
    raise ValueError(f"unsupported format {fmt!r}; expected one of {FORMATS}")


# ---------------------------------------------------------------------------------
# Templates — composition of already-fetched rows into Markdown.
# ---------------------------------------------------------------------------------
#
# Golden rule 1 lands here as an arithmetic rule: a builder prints values it was handed
# and nothing else. The one exception is a total (or a duration, or a per-row
# difference) derived from those same values, and every derived figure says on the page
# that it is derived — an unlabelled number on an invoice is indistinguishable from an
# invented one to the person reading it.

# Recharge, Billing and Invoice Approval Policy §4: the review period is fourteen days
# from issue, after which the invoice commits whether or not anyone looked at it. Stated
# on the statement itself because "silence is agreement" is the part that surprises
# people, and the statement is the last thing they read before the clock runs out.
INVOICE_REVIEW_DAYS = 14

_MISSING = "not recorded"
_EMPTY_CELL = "—"
_TWO_PLACES = Decimal("0.01")


def _text(value: object, missing: str = _MISSING) -> str:
    """Render a value for prose, naming an absence instead of printing `None`."""
    if value is None:
        return missing
    text = str(value).strip()
    return text or missing


def _cell(value: object) -> str:
    """Render a value for a table cell: one line, no pipes.

    `_blocks` splits a row on the pipe and reportlab refuses a row with more cells than
    the header, so a single pipe inside a line description turns a valid document into
    an exception at PDF time. A newline inside a cell ends the table early for the same
    class of reason.
    """
    return " ".join(_text(value, _EMPTY_CELL).replace("|", "/").split())


def _join(values: object) -> str:
    """Comma-join a list-ish value (techniques, sample types) without inventing order."""
    if values is None or isinstance(values, str):
        return _text(values, _EMPTY_CELL)
    try:
        items = [str(v).strip() for v in values if str(v).strip()]
    except TypeError:
        return _text(values, _EMPTY_CELL)
    return ", ".join(items) if items else _EMPTY_CELL


def _decimal(value: object) -> Decimal | None:
    """Parse to Decimal, or None if the value is not a number we can add up."""
    if isinstance(value, Decimal):
        return value
    if value is None or isinstance(value, bool):
        return None
    try:
        return Decimal(str(value).strip())
    except (InvalidOperation, ValueError):
        return None


def _sum(values: object) -> Decimal | None:
    """Sum exactly, or return None if any value is missing or not a number.

    Decimal rather than float because a statement whose lines visibly do not add up to
    its own total is worse than no total; None rather than a partial sum, because a
    figure quietly computed over some of the rows is a wrong number that looks right.
    """
    total = Decimal("0")
    for value in values:
        number = _decimal(value)
        if number is None:
            return None
        total += number
    return total


def _two_places(number: Decimal) -> Decimal:
    """Pad a derived figure out to the two places money and hours are written with.

    Padding only. Rounding here would make a derived total differ from the lines it was
    derived from, which is the exact discrepancy these documents exist to expose.
    """
    return number.quantize(_TWO_PLACES) if -number.as_tuple().exponent < 2 else number


def _money(value: object) -> str:
    """Format an amount without re-spelling it.

    The billing rows carry Decimal on purpose: Decimal('2689.00') through float prints
    $2689.0, which is the same number and the wrong document. Instrument hourly rates are
    NOT Decimal — they arrive as floats — so printing them verbatim reintroduced exactly
    that bug one column over ($42.0). Anything that is not already a Decimal is quantised
    to two places; a Decimal is still printed exactly as the ledger spelled it.
    """
    if value is None:
        return _MISSING
    if isinstance(value, Decimal):
        return f"${value}"
    number = _decimal(value)
    return f"${_two_places(number)}" if number is not None else f"${value}"


def _derived_money(number: Decimal | None) -> str:
    return _MISSING if number is None else _money(_two_places(number))


def _kv(pairs: list[tuple[str, str]]) -> list[str]:
    """Label/value lines using the two-space break `_blocks` turns into a real break."""
    lines = [f"{label}: {value}  " for label, value in pairs]
    if lines:
        lines[-1] = lines[-1].rstrip()
    return lines


def _subject(value: object) -> str:
    """Name a lab/user/facility given either the row a tool returned or a bare name."""
    if not isinstance(value, dict):
        return _text(value)
    name = value.get("name")
    identifier = value.get("id") or value.get("code")
    if name and identifier and str(name) != str(identifier):
        return f"{name} (`{identifier}`)"
    return _text(name or identifier)


def _place(row: dict, facility: object) -> dict:
    """Merge a row over its facility for location purposes.

    Both carry a `room` since migration 008, and the instrument's own room is the one a
    user has to walk to; the campus and building only exist on the facility.
    """
    return {**facility, **row} if isinstance(facility, dict) else row


def _location(row: dict) -> str:
    """Where something physically is, from whichever of campus/building/room exist.

    "Advanced Imaging Core" is not an answer to "where do I take the sample?", and a
    row that was never given a building must say so rather than borrow one.
    """
    place = " · ".join(str(row[key]).strip() for key in ("campus", "building") if row.get(key))
    room = row.get("room")
    if room:
        place = f"{place}, room {room}" if place else f"room {room}"
    return place or _MISSING


def _parse_dt(value: object) -> datetime | None:
    if isinstance(value, datetime):
        return value
    if not isinstance(value, str):
        return None
    try:
        return datetime.fromisoformat(value.strip().replace("Z", "+00:00"))
    except ValueError:
        return None


def _hours_between(starts: object, ends: object) -> Decimal | None:
    """Length of a window in hours, or None if either end is missing or unparseable."""
    start, end = _parse_dt(starts), _parse_dt(ends)
    if start is None or end is None:
        return None
    try:
        seconds = Decimal((end - start).total_seconds())
    except TypeError:
        # One end tz-aware and the other naive: subtracting them raises, and guessing
        # which zone the naive one meant is exactly the invention this refuses to make.
        return None
    return _two_places(seconds / Decimal(3600))


def _line_hours(line: dict) -> object:
    """Hours on a charge line, under whichever name the caller's rows use.

    reporting.v_billing_lines has no hours column at all and infinity.invoice_lines
    calls the same figure `qty`, so a caller joining either one gets the column it has
    rather than a blank.
    """
    for key in ("hours", "qty", "quantity"):
        if line.get(key) is not None:
            return line[key]
    return None


def build_invoice_statement(account_code: str, period: str, lines: list[dict],
                            total: object = None, lab: object = None) -> str:
    """Compose a statement for one account code and one period.

    Guarantee: every amount printed is one the caller passed. `total` is printed as the
    invoice's own total; the sum of the lines is derived separately and, when the two
    disagree, both appear with the derived one named as derived. A statement that
    quietly prints a total its own lines do not add up to is the one thing an invoice
    may never do, and the reader has fourteen days to query it.
    """
    rows = list(lines or [])
    # No lines and no total is silence, not zero: a statement reading $0.00 tells the
    # reader they owe nothing, which is a different claim from "nothing was supplied".
    derived = _sum(row.get("amount") for row in rows) if rows else None

    out = [
        f"# Invoice statement — {account_code}",
        "",
        *_kv([
            ("Account code", f"`{account_code}`"),
            ("Lab", _subject(lab)),
            ("Period", _text(period)),
            ("Charge lines", str(len(rows))),
        ]),
        "",
        "## Charges",
        "",
        "| Instrument | Description | Hours | Amount |",
        "|---|---|---:|---:|",
    ]
    out += [
        f"| {_cell(row.get('instrument'))} | {_cell(row.get('description'))} "
        f"| {_cell(_line_hours(row))} | {_cell(_money(row.get('amount')))} |"
        for row in rows
    ]
    if not rows:
        out.append("| _no charges in this period_ | | | |")

    out.append("")
    if total is not None:
        out.append(f"**Total for {_text(period)}: {_money(total)}**")
        if derived is not None and derived != _decimal(total):
            out += [
                "",
                f"The lines above sum to {_derived_money(derived)}, which is derived from "
                "them and does not match the invoice total. Query this with the core "
                "before the review period ends.",
            ]
    elif derived is not None:
        out.append(
            f"**Total for {_text(period)}: {_derived_money(derived)}** — derived by adding "
            "the lines above, because no invoice total was supplied."
        )
    else:
        out.append(f"**Total for {_text(period)}: {_MISSING}**")

    out += [
        "",
        "## Review and payment",
        "",
        f"- This statement covers account code `{account_code}` for the period above, and "
        "nothing else.",
        f"- The review period is {INVOICE_REVIEW_DAYS} days from the date the invoice was "
        "issued. A line can be queried, and corrected if the query is upheld, until then.",
        "- At the end of the review period the invoice commits automatically and posts to "
        "the institution's financial system, whether or not anyone looked at it: silence "
        "is agreement.",
        "- After it has posted, a correction is a credit note against a later invoice "
        "rather than an edit to this one.",
    ]
    return "\n".join(out) + "\n"


def build_booking_confirmation(booking: dict) -> str:
    """Compose the confirmation for a single booking.

    Guarantee: instrument, location, window, account code and status are printed exactly
    as the booking row gave them, and a field the row does not carry is named as not
    recorded rather than filled in. The only figure not taken verbatim is the duration,
    which is derived from the window and labelled as derived.
    """
    row = booking or {}
    facility = row.get("facility")
    facility_name = _subject(facility) if isinstance(facility, dict) else _text(facility)
    where = _location(_place(row, facility))

    duration = _hours_between(row.get("starts_at"), row.get("ends_at"))
    duration_text = (
        f"{duration} h — derived from the window above" if duration is not None else _MISSING
    )
    reference = row.get("id")

    out = [
        f"# Booking confirmation — {_text(row.get('instrument'))}",
        "",
        *_kv([
            ("Reference", f"`{reference}`" if reference else _MISSING),
            ("Instrument", _text(row.get("instrument"))),
            ("Facility", facility_name),
            ("Where", where),
            ("Starts", _text(row.get("starts_at"))),
            ("Ends", _text(row.get("ends_at"))),
            ("Duration", duration_text),
            ("Account code", f"`{row['account_code']}`" if row.get("account_code") else _MISSING),
            ("Status", f"**{_text(row.get('status'))}**"),
        ]),
        "",
        "## Changing or cancelling this booking",
        "",
    ]
    out += [
        (
            f"- Cancel or shorten it in the calendar against reference `{reference}`, before "
            "the session starts."
            if reference
            else "- Cancel or shorten it in the calendar, before the session starts."
        ),
        "- The calendar is the record. Telling a member of staff at the bench does not "
        "release the slot, and the slot stays charged to the account code above.",
        "- Moving a booking counts as a cancellation plus a new booking, so it is subject "
        "to the same notice.",
        "- The notice period and any late-cancellation charge are set by the Booking and "
        "Cancellation Rules, which this confirmation deliberately does not restate.",
    ]
    return "\n".join(out) + "\n"


def build_facility_directory(facilities: list[dict]) -> str:
    """Compose a directory of cores: where each one is, and what it has.

    Guarantee: one section per facility passed, listing only the instruments attached to
    that facility's row. A core with no instruments in the data says so — an empty list
    is a fact about the catalogue, not a prompt to go looking elsewhere.
    """
    cores = list(facilities or [])
    out = [
        "# Facility directory",
        "",
        *_kv([("Cores listed", str(len(cores)))]),
    ]
    if not cores:
        out += ["", "No facilities were supplied for this directory."]
        return "\n".join(out) + "\n"

    for core in cores:
        instruments = list(core.get("instruments") or [])
        out += [
            "",
            f"## {_text(core.get('name'))}",
            "",
            *_kv([
                ("Code", f"`{core['code']}`" if core.get("code") else _MISSING),
                ("Campus", _text(core.get("campus"))),
                ("Building", _text(core.get("building"))),
                ("Room", _text(core.get("room"))),
                ("Address", _text(core.get("address"))),
                ("Contact", _text(core.get("contact_email"))),
                ("Opening hours", _text(core.get("opening_hours"))),
            ]),
            "",
            "| Instrument | Techniques | Rate / h | Status |",
            "|---|---|---:|---|",
        ]
        out += [
            f"| {_cell(ins.get('name'))} | {_cell(_join(ins.get('techniques')))} "
            f"| {_cell(_money(ins.get('hourly_rate')))} | {_cell(ins.get('status'))} |"
            for ins in instruments
        ]
        if not instruments:
            out.append("| _no instruments listed for this core_ | | | |")
    return "\n".join(out) + "\n"


def _match_reason(row: dict) -> str:
    """Why this instrument came back, in the caller's own words."""
    # recommend_instrument names this key `why_matched`; without it every row of a
    # capability report printed "not recorded" in the one column that justifies the
    # ranking, which is the column the reader is actually there for.
    for key in ("why_matched", "matched_on", "matched_techniques", "matched_sample_types"):
        if row.get(key):
            return _join(row[key])
    for key in ("match_reason", "reason", "why"):
        if row.get(key):
            return _text(row[key])
    return _MISSING


def build_capability_report(goal: str, matches: list[dict]) -> str:
    """Compose "what can do this, and where is it" for one stated goal.

    Guarantee: the report names only instruments in `matches`, and states for each one
    why it matched using the reason the caller supplied. An empty match list produces a
    report that recommends nothing — the near miss is what makes a user book a week on
    the wrong instrument.
    """
    rows = list(matches or [])
    out = [
        f"# Capability report — {_text(goal)}",
        "",
        *_kv([
            ("Asked for", _text(goal)),
            ("Instruments matched", str(len(rows))),
        ]),
    ]

    if not rows:
        out += [
            "",
            "## No match",
            "",
            "Nothing in the catalogue matched this request. Nothing is listed here rather "
            "than the closest alternative, because a near miss reads as an answer and is "
            "the reason someone books a week on the wrong instrument. Ask the core manager "
            "whether the work can be done another way.",
        ]
        return "\n".join(out) + "\n"

    out += [
        "",
        "## Matches",
        "",
        "| Instrument | Facility | Matched on | Rate / h | Status |",
        "|---|---|---|---:|---|",
    ]
    out += [
        f"| {_cell(row.get('name') or row.get('instrument'))} "
        f"| {_cell(row.get('facility'))} | {_cell(_match_reason(row))} "
        f"| {_cell(_money(row.get('hourly_rate')))} | {_cell(row.get('status'))} |"
        for row in rows
    ]

    for position, row in enumerate(rows, start=1):
        name = _text(row.get("name") or row.get("instrument"))
        out += [
            "",
            f"### {position}. {name}",
            "",
            *_kv([
                ("Matched on", _match_reason(row)),
                ("Facility", _subject(row.get("facility"))),
                ("Where", _location(_place(row, row.get("facility")))),
                ("Modality", _text(row.get("modality"))),
                ("Techniques", _join(row.get("techniques"))),
                ("Sample types", _join(row.get("sample_types"))),
                ("Specification", _text(row.get("specification"))),
                ("Rate", f"{_money(row.get('hourly_rate'))} per hour"
                         if row.get("hourly_rate") is not None else _MISSING),
                ("Status", f"**{_text(row.get('status'))}**"),
            ]),
        ]
    return "\n".join(out) + "\n"


def build_usage_summary(user_or_lab: object, period: object, rows: list[dict],
                        totals: dict | None = None) -> str:
    """Compose scheduled versus tracked hours per instrument, with the difference.

    Guarantee: scheduled and tracked hours are printed as given. The per-row difference
    is derived (tracked minus scheduled) and its column says so; the totals row uses the
    caller's totals when it has them, and when those disagree with the rows the derived
    row sums are printed alongside. The gap between booked and used time is the whole
    point of the document, so it is never quietly reconciled.
    """
    records = list(rows or [])
    out = [
        f"# Usage summary — {_subject(user_or_lab)}",
        "",
        *_kv([
            ("Subject", _subject(user_or_lab)),
            ("Period", _text(period, "all recorded months")),
            ("Rows", str(len(records))),
        ]),
        "",
        "| Instrument | Month | Scheduled h | Tracked h | Difference h (derived) |",
        "|---|---|---:|---:|---:|",
    ]
    for record in records:
        scheduled = _decimal(record.get("scheduled_hours"))
        tracked = _decimal(record.get("tracked_hours"))
        difference = (
            _two_places(tracked - scheduled)
            if scheduled is not None and tracked is not None
            else None
        )
        out.append(
            f"| {_cell(record.get('instrument'))} | {_cell(record.get('month'))} "
            f"| {_cell(record.get('scheduled_hours'))} | {_cell(record.get('tracked_hours'))} "
            f"| {_cell(difference)} |"
        )
    if not records:
        out.append("| _no usage recorded_ | | | | |")

    summed_scheduled = _sum(r.get("scheduled_hours") for r in records)
    summed_tracked = _sum(r.get("tracked_hours") for r in records)
    given = totals or {}
    scheduled_total = given.get("scheduled_hours", summed_scheduled)
    tracked_total = given.get("tracked_hours", summed_tracked)

    scheduled_number, tracked_number = _decimal(scheduled_total), _decimal(tracked_total)
    derived_difference = (
        _two_places(tracked_number - scheduled_number)
        if scheduled_number is not None and tracked_number is not None
        else None
    )
    difference_total = given.get("difference_hours", derived_difference)

    out.append(
        f"| **Total** | | {_cell(scheduled_total)} | {_cell(tracked_total)} "
        f"| {_cell(difference_total)} |"
    )

    # A totals row that silently disagrees with the rows it sits under is how an hours
    # dispute starts, so the two are compared here rather than left to the reader.
    disagrees = [
        (label, given[key], ours)
        for label, key, ours in (
            ("scheduled", "scheduled_hours", summed_scheduled),
            ("tracked", "tracked_hours", summed_tracked),
            ("difference", "difference_hours", derived_difference),
        )
        if given.get(key) is not None and ours is not None and _decimal(given[key]) != ours
    ]
    out.append("")
    if disagrees:
        spelled = "; ".join(
            f"{label} was given as {supplied} against {_two_places(ours)}"
            for label, supplied, ours in disagrees
        )
        out.append(
            f"The totals row does not match the rows above it: {spelled}. It prints the "
            "totals as they were supplied; the second figure in each pair is the one this "
            "document's own rows give."
        )
        out.append("")
    out.append(
        "Difference is tracked hours minus scheduled hours, derived per row. A negative "
        "figure means less time was used than was booked; the booked time is what was "
        "charged, so a persistent gap is worth raising with the core."
    )
    return "\n".join(out) + "\n"
